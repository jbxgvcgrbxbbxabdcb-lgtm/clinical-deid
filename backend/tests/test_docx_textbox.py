"""DOCX text-box extraction and write-back."""

from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from lxml import etree

from backend.deid.docx_ooxml import (
    W_NS,
    apply_textbox_redactions,
    load_enriched_docx,
)
from backend.deid.ops import apply_selected_docx_redaction, run_docx_review
from backend.deid.rules import resolve_custom_recognizer


def _make_textbox_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Body contact: body@example.com")
    paragraph = doc.add_paragraph()
    pict_xml = """
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:v="urn:schemas-microsoft-com:vml"
         xmlns:o="urn:schemas-microsoft-com:office:office">
      <w:pict>
        <v:shapetype id="_x0000_t202" coordsize="21600,21600" o:spt="202"
                     path="m,l,21600r21600,l21600,xe">
          <v:stroke joinstyle="miter"/>
          <v:path gradientshapeok="t" o:connecttype="rect"/>
        </v:shapetype>
        <v:shape id="TextBox 1" type="#_x0000_t202"
                 style="width:200pt;height:50pt" filled="t" stroked="t">
          <v:textbox>
            <w:txbxContent>
              <w:p>
                <w:r>
                  <w:t>Box secret: textbox@example.com</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </v:textbox>
        </v:shape>
      </w:pict>
    </w:r>
    """
    paragraph._p.append(etree.fromstring(pict_xml))
    doc.save(path)


def _textbox_plain(path: Path) -> str:
    with ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    texts = []
    for box in root.iter(f"{{{W_NS}}}txbxContent"):
        for node in box.iter(f"{{{W_NS}}}t"):
            texts.append(node.text or "")
    return "".join(texts)


def test_enrichment_includes_textbox_missed_by_openmed(tmp_path: Path) -> None:
    source = tmp_path / "note.docx"
    _make_textbox_docx(source)
    enriched = load_enriched_docx(source)
    assert "body@example.com" in enriched.text
    assert "textbox@example.com" in enriched.text
    assert enriched.base_length < len(enriched.text)
    assert len(enriched.textboxes) == 1


def test_textbox_redaction_rewrites_ooxml(tmp_path: Path) -> None:
    source = tmp_path / "note.docx"
    output = tmp_path / "out.docx"
    _make_textbox_docx(source)
    enriched = load_enriched_docx(source)
    region = enriched.textboxes[0]
    needle = "textbox@example.com"
    local = region.text.index(needle)
    apply_textbox_redactions(
        source,
        output,
        enriched.textboxes,
        [
            {
                "start": region.start + local,
                "end": region.start + local + len(needle),
                "replacement": "[email]",
            }
        ],
    )
    assert "textbox@example.com" not in _textbox_plain(output)
    assert "[email]" in _textbox_plain(output)


def test_partition_handles_review_entity_at_offset_zero() -> None:
    from backend.deid.docx_ooxml import EnrichedDocx, partition_docx_spans
    from backend.deid.models import ReviewEntity

    enriched = EnrichedDocx(text="ab\nbox", base_length=2, textboxes=())
    entity = ReviewEntity(
        id="e0",
        label="email",
        text="a",
        start=0,
        end=1,
        confidence=1.0,
        replacement="[email]",
    )
    body, boxes = partition_docx_spans(enriched, [entity])
    assert body == [entity]
    assert boxes == []


def test_docx_review_and_apply_cover_textbox(tmp_path: Path) -> None:
    source = tmp_path / "note.docx"
    _make_textbox_docx(source)
    recognizer = resolve_custom_recognizer()
    view = run_docx_review(source, "mask", custom_recognizer=recognizer)
    texts = {entity.text for entity in view.entities}
    assert "textbox@example.com" in texts
    assert "body@example.com" in texts

    selected = [
        {
            "id": entity.id,
            "label": entity.label,
            "text": entity.text,
            "start": entity.start,
            "end": entity.end,
            "confidence": entity.confidence,
            "replacement": entity.replacement,
        }
        for entity in view.entities
        if "example.com" in entity.text
    ]
    result = apply_selected_docx_redaction(
        source,
        "mask",
        selected,
        custom_recognizer=recognizer,
        review_entities=view.entities,
    )
    assert result.output_path is not None
    out = Path(result.output_path)
    plain_box = _textbox_plain(out)
    assert "textbox@example.com" not in plain_box
    assert "textbox@example.com" not in result.deidentified_text
    assert "body@example.com" not in result.deidentified_text
